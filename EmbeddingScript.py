import os
import django
import numpy as np
import json

os.environ.setdefault("DJANGO_SETTINGS_MODULE", "chatBot.settings")
django.setup()

from app1.llm_service import query
from docx import Document
import mysql.connector
from mysql.connector import Error

# Function to read and create chunks
def extract_chunks_by_headings(doc, headings):
    chunks = {}
    current_heading = None
    current_text = []

    for paragraph in doc.paragraphs:
        # Check if the paragraph is a heading 
        if paragraph.text.strip() in headings:
            if current_heading and current_text:
                # Save the previous chunk
                chunks[current_heading] = "\n".join(current_text)
                current_text = []
            # Start a new chunk for the current heading
            current_heading = paragraph.text.strip()
        elif current_heading:
            # Append text to the current chunk
            current_text.append(paragraph.text.strip())

    # Save the last chunk
    if current_heading and current_text:
        chunks[current_heading] = "\n".join(current_text)

    return chunks


# Function to insert embeddings into MySQL
def insert_embeddings_into_mysql(headings, embeddings):
    try:
        db = mysql.connector.connect(
            host="localhost",
            user="root",
            passwd="",
            database="foachatbot"
        )
        cursor = db.cursor()

    
        create_table_query = """
        CREATE TABLE IF NOT EXISTS embeddings (
            id INT AUTO_INCREMENT PRIMARY KEY,
            heading VARCHAR(255) NOT NULL,
            embedding JSON NOT NULL,
            content TEXT NOT NULL
        )
        """
        cursor.execute(create_table_query)
        insert_query = "INSERT INTO embeddings (heading, embedding, content) VALUES (%s, %s, %s)"
        for heading, embedding, content in zip(headings, embeddings, documents):
            embedding_list = embedding.tolist() if isinstance(embedding, np.ndarray) else embedding
            cursor.execute(insert_query, (heading, json.dumps(embedding_list), content))

        db.commit()
        print("Embeddings inserted successfully!")

    except Error as e:
        print(f"Error: {e}")
    finally:
        if db.is_connected():
            cursor.close()
            db.close()




# doc = Document('/Users/apple/Downloads/RAG/Dataset-1.docx')
doc = Document(r'D:\Python Project\RAG-Chatbot - Copy\Dataset-2.docx')


# Define the headings you want to extract
headings = [
    "Purpose:",
    "Teaching & Learning Environment",
    "Student Progress Tracking",
    "Technical Requirements",
    "Certification & Course Completion",
    "Fee Structure & Discounts",
    "Timing Flexibility",
    "Free Courses",
    "Course Catalog",
    "Key Areas Covered:",
    "Quran Reading Course (Nazra)",
    "Course Outline:",
    "Handling Specific Requests",
    "Contact Information"
    ]

# Extract chunks based on headings
chunks = extract_chunks_by_headings(doc, headings)
# Prepare documents for embedding
documents = list(chunks.values())  # Extract the content (values) from the chunks dictionary
# Get embeddings for the documents
dataEmbeddings = query(documents)
# Insert embeddings into MySQL
insert_embeddings_into_mysql(list(chunks.keys()), dataEmbeddings)




